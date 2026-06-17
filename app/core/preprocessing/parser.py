"""
Document Parser

Multi-format document parsing:
- PDF (via PyMuPDF) - preserves structure, pages, tables
- DOCX (via python-docx) - preserves headings, paragraphs
- TXT - simple text parsing
- HTML (via BeautifulSoup) - preserves structure
"""

import re
from pathlib import Path
from typing import Any

import fitz  # PyMuPDF
from docx import Document
from bs4 import BeautifulSoup

from app.config import settings
from app.utils.exceptions import DocumentProcessingError
from app.utils.logging import get_logger

logger = get_logger("core.preprocessing.parser")


class ParsedDocument:
    """Parsed document with structured content."""

    def __init__(
        self,
        filename: str,
        file_type: str,
        total_pages: int = 1,
    ):
        self.filename = filename
        self.file_type = file_type
        self.total_pages = total_pages
        self.sections: list[dict[str, Any]] = []
        self.metadata: dict[str, Any] = {}

    def add_section(
        self,
        content: str,
        page_number: int | None = None,
        section_type: str = "text",
        heading_level: int | None = None,
        heading_text: str | None = None,
        position: dict[str, Any] | None = None,
    ) -> None:
        """Add a parsed section."""
        self.sections.append({
            "content": content,
            "page_number": page_number,
            "section_type": section_type,
            "heading_level": heading_level,
            "heading_text": heading_text,
            "position": position or {},
        })

    def get_total_chars(self) -> int:
        """Get total character count."""
        return sum(len(s["content"]) for s in self.sections)

    def to_dict(self) -> dict[str, Any]:
        """Convert to dictionary."""
        return {
            "filename": self.filename,
            "file_type": self.file_type,
            "total_pages": self.total_pages,
            "total_sections": len(self.sections),
            "total_chars": self.get_total_chars(),
            "metadata": self.metadata,
            "sections": self.sections,
        }


class DocumentParser:
    """Multi-format document parser."""

    def __init__(self):
        self.supported_types = set(settings.ALLOWED_EXTENSIONS)

    def parse(
        self,
        file_path: str,
        file_type: str,
    ) -> ParsedDocument:
        """
        Parse document based on file type.

        Args:
            file_path: Path to file
            file_type: File extension (pdf, docx, txt, html)

        Returns:
            ParsedDocument with structured content
        """
        if file_type not in self.supported_types:
            raise DocumentProcessingError(
                Path(file_path).name,
                "parse",
                f"Unsupported file type: {file_type}"
            )

        filename = Path(file_path).name

        try:
            if file_type == "pdf":
                return self._parse_pdf(file_path, filename)
            elif file_type == "docx":
                return self._parse_docx(file_path, filename)
            elif file_type == "txt":
                return self._parse_txt(file_path, filename)
            elif file_type in ("html", "md"):
                return self._parse_html(file_path, filename)
            else:
                raise DocumentProcessingError(
                    filename,
                    "parse",
                    f"Parser not implemented for {file_type}"
                )

        except Exception as e:
            logger.error(f"Failed to parse {filename}: {e}")
            raise DocumentProcessingError(filename, "parse", str(e))

    def _parse_pdf(
        self,
        file_path: str,
        filename: str,
    ) -> ParsedDocument:
        """
        Parse PDF using PyMuPDF.

        Preserves:
        - Page numbers
        - Text blocks with position
        - Tables (extracted as text)
        - Headings (detected by font size)
        """
        doc = fitz.open(file_path)
        parsed = ParsedDocument(
            filename=filename,
            file_type="pdf",
            total_pages=doc.page_count,
        )

        parsed.metadata = {
            "title": doc.metadata.get("title", ""),
            "author": doc.metadata.get("author", ""),
            "subject": doc.metadata.get("subject", ""),
            "keywords": doc.metadata.get("keywords", ""),
            "creator": doc.metadata.get("creator", ""),
        }

        for page_num in range(doc.page_count):
            page = doc[page_num]

            # Get text blocks with detailed info
            blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]

            for block in blocks:
                if block.get("type") == 0:  # Text block
                    # Extract text from lines
                    text_lines = []
                    for line in block.get("lines", []):
                        line_text = "".join(span.get("text", "") for span in line.get("spans", []))
                        text_lines.append(line_text)

                    content = "\n".join(text_lines).strip()

                    if content:
                        # Detect heading by font size
                        heading_level = None
                        heading_text = None

                        # Check if first span is large font (heading)
                        first_span = None
                        for line in block.get("lines", []):
                            if line.get("spans"):
                                first_span = line["spans"][0]
                                break

                        if first_span:
                            font_size = first_span.get("size", 0)
                            # Larger font sizes indicate headings
                            if font_size > 14:
                                heading_level = 1
                                heading_text = content[:100]
                            elif font_size > 12:
                                heading_level = 2

                        parsed.add_section(
                            content=content,
                            page_number=page_num + 1,
                            section_type="text",
                            heading_level=heading_level,
                            heading_text=heading_text,
                            position={
                                "bbox": block.get("bbox", []),
                                "page_width": page.rect.width,
                                "page_height": page.rect.height,
                            },
                        )

                elif block.get("type") == 1:  # Image block - skip
                    pass

            # Try to extract tables
            tables = page.find_tables()
            if tables:
                for table in tables:
                    table_text = table.to_markdown()
                    parsed.add_section(
                        content=table_text,
                        page_number=page_num + 1,
                        section_type="table",
                    )

        doc.close()
        logger.info(f"Parsed PDF {filename}: {parsed.total_pages} pages, {len(parsed.sections)} sections")
        return parsed

    def _parse_docx(
        self,
        file_path: str,
        filename: str,
    ) -> ParsedDocument:
        """
        Parse DOCX using python-docx.

        Preserves:
        - Paragraphs with heading levels
        - Tables
        - Lists
        """
        doc = Document(file_path)
        parsed = ParsedDocument(
            filename=filename,
            file_type="docx",
            total_pages=1,  # DOCX doesn't have clear page concept
        )

        # Extract core properties
        core_props = doc.core_properties
        parsed.metadata = {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "keywords": core_props.keywords or "",
        }

        # Parse paragraphs
        for para in doc.paragraphs:
            content = para.text.strip()
            if not content:
                continue

            # Detect heading level
            heading_level = None
            heading_text = None

            if para.style.name.startswith("Heading"):
                try:
                    heading_level = int(para.style.name.split()[-1])
                    heading_text = content
                except ValueError:
                    pass
            elif para.style.name == "Title":
                heading_level = 1
                heading_text = content

            section_type = "paragraph"
            if heading_level:
                section_type = "heading"

            parsed.add_section(
                content=content,
                page_number=None,
                section_type=section_type,
                heading_level=heading_level,
                heading_text=heading_text,
            )

        # Parse tables
        for table in doc.tables:
            table_text = self._extract_table_text(table)
            parsed.add_section(
                content=table_text,
                page_number=None,
                section_type="table",
            )

        logger.info(f"Parsed DOCX {filename}: {len(parsed.sections)} sections")
        return parsed

    def _extract_table_text(self, table: Any) -> str:
        """Convert table to text format."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
        return "\n".join(rows)

    def _parse_txt(
        self,
        file_path: str,
        filename: str,
    ) -> ParsedDocument:
        """
        Parse plain text file.

        Detects:
        - Paragraphs (split by double newline)
        - Potential headings (lines ending with colon or starting with caps)
        """
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()

        parsed = ParsedDocument(
            filename=filename,
            file_type="txt",
            total_pages=1,
        )

        # Split by paragraphs
        paragraphs = re.split(r"\n\s*\n", content)

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            # Detect potential heading
            heading_level = None
            heading_text = None

            # Lines ending with colon might be headings
            if para.endswith(":") and len(para) < 100:
                heading_level = 2
                heading_text = para
            # Lines that are all caps and short might be headings
            elif para.isupper() and len(para) < 50:
                heading_level = 1
                heading_text = para

            parsed.add_section(
                content=para,
                page_number=1,
                section_type="text",
                heading_level=heading_level,
                heading_text=heading_text,
            )

        logger.info(f"Parsed TXT {filename}: {len(parsed.sections)} sections")
        return parsed

    def _parse_html(
        self,
        file_path: str,
        filename: str,
    ) -> ParsedDocument:
        """
        Parse HTML using BeautifulSoup.

        Preserves:
        - Headings (h1-h6)
        - Paragraphs
        - Lists
        - Tables
        """
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            html_content = f.read()

        soup = BeautifulSoup(html_content, "lxml")

        parsed = ParsedDocument(
            filename=filename,
            file_type="html",
            total_pages=1,
        )

        # Extract title
        title = soup.find("title")
        if title:
            parsed.metadata["title"] = title.get_text().strip()

        # Parse headings
        for level in range(1, 7):
            for heading in soup.find_all(f"h{level}"):
                content = heading.get_text().strip()
                if content:
                    parsed.add_section(
                        content=content,
                        page_number=None,
                        section_type="heading",
                        heading_level=level,
                        heading_text=content,
                    )

        # Parse paragraphs
        for para in soup.find_all("p"):
            content = para.get_text().strip()
            if content:
                parsed.add_section(
                    content=content,
                    page_number=None,
                    section_type="paragraph",
                )

        # Parse tables
        for table in soup.find_all("table"):
            table_text = self._extract_html_table(table)
            parsed.add_section(
                content=table_text,
                page_number=None,
                section_type="table",
            )

        # Parse lists
        for list_elem in soup.find_all(["ul", "ol"]):
            list_text = self._extract_list(list_elem)
            parsed.add_section(
                content=list_text,
                page_number=None,
                section_type="list",
            )

        logger.info(f"Parsed HTML {filename}: {len(parsed.sections)} sections")
        return parsed

    def _extract_html_table(self, table: Any) -> str:
        """Convert HTML table to text format."""
        rows = []
        for tr in table.find_all("tr"):
            cells = [td.get_text().strip() for td in tr.find_all(["td", "th"])]
            if cells:
                rows.append("| " + " | ".join(cells) + " |")
        return "\n".join(rows)

    def _extract_list(self, list_elem: Any) -> str:
        """Convert HTML list to text format."""
        items = []
        for li in list_elem.find_all("li"):
            text = li.get_text().strip()
            if text:
                items.append(f"- {text}")
        return "\n".join(items)


def get_parser() -> DocumentParser:
    """Get document parser instance."""
    return DocumentParser()